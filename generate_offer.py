#!/usr/bin/env python3
# -*- coding: utf-8 -*-
"""
LaudaMed — Generator Oferty Handlowej
Generates a professional A4 Portrait PDF sales offer (4 pages)
and optionally a DOCX template.

Usage:
    python3 generate_offer.py --data offer_data.json   # generate PDF
    python3 generate_offer.py --template               # generate DOCX template
    python3 generate_offer.py --demo                   # generate demo PDF + DOCX for ViV60
"""

import os
import sys
import json
import math
import argparse
from datetime import datetime, timedelta

# ── ReportLab ─────────────────────────────────────────────────────────────────
from reportlab.lib.pagesizes import A4
from reportlab.lib.units import mm
from reportlab.lib.colors import HexColor, white, black, Color
from reportlab.pdfgen import canvas
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont
from reportlab.lib.utils import ImageReader

# ── python-docx ───────────────────────────────────────────────────────────────
try:
    from docx import Document
    from docx.shared import Pt, Mm, RGBColor, Inches
    from docx.enum.text import WD_ALIGN_PARAGRAPH
    from docx.enum.table import WD_ALIGN_VERTICAL, WD_TABLE_ALIGNMENT
    from docx.oxml.ns import qn
    from docx.oxml import OxmlElement
    DOCX_AVAILABLE = True
except ImportError:
    DOCX_AVAILABLE = False

# ══════════════════════════════════════════════════════════════════════════════
# CONSTANTS & DESIGN TOKENS
# ══════════════════════════════════════════════════════════════════════════════

FONTS_DIR   = "/home/user/workspace/fonts"
IMG_DIR     = "/home/user/workspace/extracted_images"
LOGO_COLOR  = "/home/user/workspace/logo_laudamed_color_cropped.png"
LOGO_WHITE  = "/home/user/workspace/logo_laudamed_white_transparent.png"
OUTPUT_DIR  = "/home/user/workspace/oferta"

# A4 Portrait
PAGE_W, PAGE_H = A4   # 595.276 x 841.890 pts
MARGIN_L = MARGIN_R = 18 * mm
MARGIN_T = 16 * mm
MARGIN_B = 18 * mm
CONTENT_W = PAGE_W - MARGIN_L - MARGIN_R

# Brand colors
C_DARK    = HexColor('#0D2137')
C_MID     = HexColor('#1565A0')
C_LIGHT   = HexColor('#2E8BC0')
C_TEXT    = HexColor('#1A1A1A')
C_LABEL   = HexColor('#333333')
C_CAPTION = HexColor('#666666')
C_PANEL   = HexColor('#F5F8FA')
C_LINE    = HexColor('#E8EDF2')
C_WHITE   = white
C_GREEN   = HexColor('#2E7D32')

# Category accent colors
CAT_COLORS = {
    'USG':     HexColor('#1565A0'),
    'EKG':     HexColor('#2E7D32'),
    'DIAG':    HexColor('#6A1B9A'),
    'PERINAT': HexColor('#00838F'),
    'MONITOR': HexColor('#00838F'),
    'WYPOSAZ': HexColor('#37474F'),
    'OSWIETL': HexColor('#455A64'),
    'HIGIENA': HexColor('#2E7D32'),
}

# Company info
COMPANY = {
    'name':    'LaudaMed Sp. z o.o.',
    'address': 'ul. Legnicka 57D/B/D, 54-203 Wrocław',
    'tel':     '509 168 241 / 510 669 463',
    'email':   'biuro@laudamed.pl',
    'web':     'www.laudamed.pl',
}

# ══════════════════════════════════════════════════════════════════════════════
# FONT REGISTRATION
# ══════════════════════════════════════════════════════════════════════════════

def register_fonts():
    fonts = {
        'Inter':          'Inter-Regular.ttf',
        'Inter-Bold':     'Inter-Bold.ttf',
        'Inter-SemiBold': 'Inter-SemiBold.ttf',
        'Inter-Light':    'Inter-Light.ttf',
        'Inter-Medium':   'Inter-Medium.ttf',
    }
    for name, fname in fonts.items():
        path = os.path.join(FONTS_DIR, fname)
        if os.path.exists(path):
            try:
                pdfmetrics.registerFont(TTFont(name, path))
            except Exception:
                pass

# ══════════════════════════════════════════════════════════════════════════════
# HELPER DRAWING UTILITIES
# ══════════════════════════════════════════════════════════════════════════════

def hex_to_rgb01(color: HexColor):
    """Return (r,g,b) floats 0-1 for canvas.setFillColor via r,g,b."""
    return color

def draw_rect(c, x, y, w, h, fill=None, stroke=None, stroke_w=0.5, radius=0):
    c.saveState()
    if fill:
        c.setFillColor(fill)
    if stroke:
        c.setStrokeColor(stroke)
        c.setLineWidth(stroke_w)
    if radius > 0:
        c.roundRect(x, y, w, h, radius, fill=1 if fill else 0, stroke=1 if stroke else 0)
    else:
        c.rect(x, y, w, h, fill=1 if fill else 0, stroke=1 if stroke else 0)
    c.restoreState()

def draw_line(c, x1, y1, x2, y2, color=None, width=0.5):
    c.saveState()
    if color:
        c.setStrokeColor(color)
    c.setLineWidth(width)
    c.line(x1, y1, x2, y2)
    c.restoreState()

def draw_text(c, text, x, y, font='Inter', size=10, color=C_TEXT, align='left'):
    c.saveState()
    c.setFont(font, size)
    c.setFillColor(color)
    if align == 'center':
        c.drawCentredString(x, y, text)
    elif align == 'right':
        c.drawRightString(x, y, text)
    else:
        c.drawString(x, y, text)
    c.restoreState()

def draw_wrapped_text(c, text, x, y, max_w, font='Inter', size=9, color=C_TEXT,
                      line_height=None, align='left'):
    """Draw text with word-wrap. Returns final y after drawing."""
    if line_height is None:
        line_height = size * 1.45
    words = text.split()
    lines = []
    current = ''
    c.saveState()
    c.setFont(font, size)
    for word in words:
        test = (current + ' ' + word).strip()
        if c.stringWidth(test, font, size) <= max_w:
            current = test
        else:
            if current:
                lines.append(current)
            current = word
    if current:
        lines.append(current)

    c.setFillColor(color)
    for line in lines:
        if align == 'center':
            c.drawCentredString(x + max_w / 2, y, line)
        elif align == 'right':
            c.drawRightString(x + max_w, y, line)
        else:
            c.drawString(x, y, line)
        y -= line_height
    c.restoreState()
    return y

def draw_logo(c, path, x, y, w, h=None):
    """Draw logo image, preserving aspect ratio if h is None."""
    if not os.path.exists(path):
        return
    try:
        img = ImageReader(path)
        iw, ih = img.getSize()
        if h is None:
            h = w * ih / iw
        c.drawImage(img, x, y, width=w, height=h, preserveAspectRatio=True, mask='auto')
    except Exception as e:
        print(f"[warn] Logo load failed: {e}")

def draw_product_image(c, path, x, y, max_w, max_h):
    """Draw product image centered in bounding box."""
    if not path or not os.path.exists(path):
        # Draw placeholder
        draw_rect(c, x, y, max_w, max_h, fill=C_PANEL, stroke=C_LINE)
        draw_text(c, 'Zdjęcie produktu', x + max_w/2, y + max_h/2, 
                  font='Inter', size=9, color=C_CAPTION, align='center')
        return
    try:
        img = ImageReader(path)
        iw, ih = img.getSize()
        scale = min(max_w / iw, max_h / ih)
        dw, dh = iw * scale, ih * scale
        dx = x + (max_w - dw) / 2
        dy = y + (max_h - dh) / 2
        c.drawImage(img, dx, dy, width=dw, height=dh, mask='auto')
    except Exception as e:
        print(f"[warn] Product image load failed: {e}")

def pill_badge(c, x, y, label, bg_color, text_color=white, font='Inter-SemiBold', size=7.5):
    """Draw a pill-shaped badge."""
    c.saveState()
    c.setFont(font, size)
    tw = c.stringWidth(label, font, size)
    pad_h, pad_v = 5, 3
    w = tw + pad_h * 2
    h = size + pad_v * 2
    r = h / 2
    c.setFillColor(bg_color)
    c.roundRect(x, y, w, h, r, fill=1, stroke=0)
    c.setFillColor(text_color)
    c.drawString(x + pad_h, y + pad_v + 0.5, label)
    c.restoreState()
    return w   # return width for chaining

def format_price(value: float) -> str:
    """Format price with space thousands separator."""
    return f"{value:,.0f}".replace(',', ' ')

def section_title(c, x, y, text, color=C_MID, width=None):
    """Draw a section title with colored left bar and line."""
    bar_w = 3
    bar_h = 13
    draw_rect(c, x, y - bar_h + 3, bar_w, bar_h, fill=color)
    draw_text(c, text, x + bar_w + 5, y, font='Inter-Bold', size=10.5, color=C_DARK)
    if width:
        draw_line(c, x, y - bar_h + 1, x + width, y - bar_h + 1, color=C_LINE)
    return y - bar_h - 3

# ══════════════════════════════════════════════════════════════════════════════
# HEADER & FOOTER
# ══════════════════════════════════════════════════════════════════════════════

HEADER_H = 14 * mm
FOOTER_H = 10 * mm

def draw_header(c, offer_data):
    """Draw header bar for pages 2-4."""
    accent = CAT_COLORS.get(offer_data.get('cat_key', 'USG'), C_MID)
    
    # Dark background bar
    draw_rect(c, 0, PAGE_H - HEADER_H, PAGE_W, HEADER_H, fill=C_DARK)
    
    # Colored accent bar — thin bottom strip
    draw_rect(c, 0, PAGE_H - HEADER_H, PAGE_W, 2, fill=accent)
    
    # Logo (white)
    logo_w = 36 * mm
    logo_h = HEADER_H * 0.6
    logo_y = PAGE_H - HEADER_H + (HEADER_H - logo_h) / 2
    draw_logo(c, LOGO_WHITE, MARGIN_L, logo_y, logo_w, logo_h)
    
    # Right side: offer label
    label = f"OFERTA HANDLOWA | {offer_data.get('date','')[:4]}"
    draw_text(c, label, PAGE_W - MARGIN_R, PAGE_H - HEADER_H/2 - 2,
              font='Inter-SemiBold', size=8, color=C_WHITE, align='right')
    
    # Offer number
    draw_text(c, offer_data.get('offer_number', ''), 
              PAGE_W - MARGIN_R, PAGE_H - HEADER_H/2 + 5,
              font='Inter', size=7, color=HexColor('#AACCEE'), align='right')

def draw_footer(c, page_num=None, total_pages=None):
    """Draw footer bar."""
    y = MARGIN_B
    # Line
    draw_line(c, MARGIN_L, y + FOOTER_H - 1, PAGE_W - MARGIN_R, y + FOOTER_H - 1,
              color=C_LINE)
    
    # Left: company info
    info = f"{COMPANY['name']}  -  {COMPANY['address']}"
    draw_text(c, info, MARGIN_L, y + FOOTER_H/2, 
              font='Inter', size=6.5, color=C_CAPTION)
    
    # Center: contacts
    contact = f"Tel: {COMPANY['tel']}  /  {COMPANY['email']}  /  {COMPANY['web']}"
    draw_text(c, contact, PAGE_W/2, y + FOOTER_H/2,
              font='Inter', size=6.5, color=C_CAPTION, align='center')
    
    # Right: page number
    if page_num:
        ptext = f"{page_num}"
        if total_pages:
            ptext += f" / {total_pages}"
        draw_text(c, ptext, PAGE_W - MARGIN_R, y + FOOTER_H/2,
                  font='Inter', size=7, color=C_CAPTION, align='right')

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 1 — COVER
# ══════════════════════════════════════════════════════════════════════════════

def build_cover_page(c, offer_data):
    c.saveState()
    
    # ── Dark gradient background (simulated with rectangles) ──────────────────
    steps = 60
    for i in range(steps):
        frac = i / steps
        r = 0x0D + (0x15 - 0x0D) * frac
        g = 0x21 + (0x65 - 0x21) * frac
        b = 0x37 + (0xA0 - 0x37) * frac
        color = Color(r/255, g/255, b/255)
        strip_h = PAGE_H / steps
        draw_rect(c, 0, PAGE_H - (i+1)*strip_h, PAGE_W, strip_h + 1, fill=color)
    
    # ── Decorative elements ───────────────────────────────────────────────────
    accent = CAT_COLORS.get(offer_data.get('cat_key', 'USG'), C_LIGHT)
    
    # Top accent line
    draw_rect(c, 0, PAGE_H - 3, PAGE_W, 3, fill=accent)
    
    # Subtle diagonal lines (decorative)
    c.saveState()
    c.setStrokeColor(Color(1, 1, 1, 0.04))
    c.setLineWidth(40)
    for xi in range(-2, 4):
        c.line(PAGE_W * xi * 0.3, 0, PAGE_W * xi * 0.3 + PAGE_H, PAGE_H)
    c.restoreState()
    
    # ── HERO IMAGE (optional, upper right area) ───────────────────────────────
    hero_path = offer_data.get('hero_image') or offer_data.get('product_image')
    if hero_path and os.path.exists(hero_path):
        try:
            img = ImageReader(hero_path)
            iw, ih = img.getSize()
            hero_w = PAGE_W * 0.5
            hero_h = hero_w * ih / iw
            hero_h = min(hero_h, PAGE_H * 0.45)
            hero_w = hero_h * iw / ih
            # Center horizontally
            hx = (PAGE_W - hero_w) / 2
            hy = PAGE_H * 0.38
            # Fade mask / simple alpha draw
            c.saveState()
            c.setFillAlpha(0.9)
            c.drawImage(img, hx, hy, width=hero_w, height=hero_h, 
                       mask='auto', preserveAspectRatio=True)
            c.restoreState()
        except Exception:
            pass
    
    # ── LOGO (white, centered top) ────────────────────────────────────────────
    logo_w = 65 * mm
    lx = (PAGE_W - logo_w) / 2
    ly = PAGE_H - 42 * mm
    draw_logo(c, LOGO_WHITE, lx, ly, logo_w)
    
    # ── "OFERTA HANDLOWA" label ───────────────────────────────────────────────
    label_y = PAGE_H - 58 * mm
    label = "OFERTA HANDLOWA"
    draw_rect(c, 0, label_y - 5, PAGE_W, 18, fill=Color(1,1,1,0.08))
    draw_text(c, label, PAGE_W/2, label_y, font='Inter-SemiBold', size=9,
              color=Color(0.8, 0.88, 1, 1), align='center')
    
    # ── PRODUCT NAME ──────────────────────────────────────────────────────────
    prod_name = offer_data.get('product_name', '')
    prod_model = offer_data.get('product_model', '')
    category = offer_data.get('product_category', '')
    manufacturer = offer_data.get('manufacturer', '')
    
    # Center block
    center_y = PAGE_H * 0.42
    
    # Category + manufacturer
    cat_text = f"{category}  ·  {manufacturer}"
    draw_text(c, cat_text, PAGE_W/2, center_y + 60,
              font='Inter-Light', size=10, color=Color(0.65, 0.8, 1, 1), align='center')
    
    # Product name — large
    name_size = 28
    c.saveState()
    c.setFont('Inter-Bold', name_size)
    name_w = c.stringWidth(prod_name, 'Inter-Bold', name_size)
    while name_w > PAGE_W - 40*mm and name_size > 18:
        name_size -= 1
        name_w = c.stringWidth(prod_name, 'Inter-Bold', name_size)
    c.restoreState()
    draw_text(c, prod_name, PAGE_W/2, center_y + 32,
              font='Inter-Bold', size=name_size, color=white, align='center')
    
    # Model / subtitle
    if prod_model and prod_model != prod_name:
        draw_text(c, f"Model: {prod_model}", PAGE_W/2, center_y + 10,
                  font='Inter', size=11, color=Color(0.75, 0.88, 1, 1), align='center')
    
    # Accent separator line
    line_y = center_y - 6
    line_w = 80 * mm
    draw_line(c, (PAGE_W - line_w)/2, line_y, (PAGE_W + line_w)/2, line_y,
              color=accent, width=1.5)
    
    # ── CERT BADGES ──────────────────────────────────────────────────────────
    certs = offer_data.get('certs', [])
    if certs:
        badge_y = center_y - 22
        total_badge_w = len(certs) * 36 + (len(certs)-1) * 6
        bx = (PAGE_W - total_badge_w) / 2
        for cert in certs:
            c.saveState()
            c.setFillColor(Color(1,1,1,0.12))
            c.setStrokeColor(Color(1,1,1,0.3))
            c.setLineWidth(0.5)
            c.roundRect(bx, badge_y - 2, 36, 13, 6, fill=1, stroke=1)
            c.setFont('Inter-SemiBold', 7)
            c.setFillColor(white)
            c.drawCentredString(bx + 18, badge_y + 2, cert)
            c.restoreState()
            bx += 42
    
    # ── OFFER INFO BOX ────────────────────────────────────────────────────────
    box_y = 95  # from bottom
    box_h = 40
    draw_rect(c, MARGIN_L, box_y, CONTENT_W, box_h, fill=Color(0,0,0,0.25))
    draw_line(c, MARGIN_L, box_y + box_h, MARGIN_L + CONTENT_W, box_y + box_h,
              color=accent, width=1)
    
    # Offer number
    draw_text(c, 'Nr oferty:', MARGIN_L + 8, box_y + 26,
              font='Inter-Light', size=7.5, color=Color(0.7,0.82,1,1))
    draw_text(c, offer_data.get('offer_number', '—'), MARGIN_L + 8, box_y + 15,
              font='Inter-SemiBold', size=9, color=white)
    
    # Date
    cx = PAGE_W / 2
    draw_text(c, 'Data oferty:', cx, box_y + 26,
              font='Inter-Light', size=7.5, color=Color(0.7,0.82,1,1))
    draw_text(c, offer_data.get('date', '—'), cx, box_y + 15,
              font='Inter-SemiBold', size=9, color=white)
    
    # Valid until
    rx = PAGE_W - MARGIN_R
    draw_text(c, 'Ważna do:', rx - 85, box_y + 26,
              font='Inter-Light', size=7.5, color=Color(0.7,0.82,1,1))
    draw_text(c, offer_data.get('valid_until', '—'), rx - 85, box_y + 15,
              font='Inter-SemiBold', size=9, color=white)
    
    # Client (if given)
    client = offer_data.get('client_name', '')
    if client:
        draw_text(c, 'Oferta przygotowana dla:', MARGIN_L + 8, box_y + 38,
                  font='Inter-Light', size=7, color=Color(0.65,0.78,1,1))
        draw_text(c, client, MARGIN_L + 8, box_y + 28,
                  font='Inter-SemiBold', size=8.5, color=white)
    
    # ── BOTTOM CONTACT STRIP ─────────────────────────────────────────────────
    strip_h = 28
    draw_rect(c, 0, 0, PAGE_W, strip_h, fill=Color(0,0,0,0.4))
    
    contact_str = (f"{COMPANY['name']}  |  {COMPANY['address']}  |  "
                   f"Tel: {COMPANY['tel']}  |  {COMPANY['email']}  |  {COMPANY['web']}")
    draw_text(c, contact_str, PAGE_W/2, 10, font='Inter', size=7,
              color=Color(0.75, 0.85, 1, 1), align='center')
    
    draw_line(c, 0, strip_h, PAGE_W, strip_h, color=Color(1,1,1,0.1), width=0.5)
    
    c.restoreState()
    c.showPage()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 2 — PRODUCT CARD
# ══════════════════════════════════════════════════════════════════════════════

def build_product_card_page(c, offer_data):
    draw_header(c, offer_data)
    draw_footer(c, page_num=2, total_pages=4)
    
    accent = CAT_COLORS.get(offer_data.get('cat_key', 'USG'), C_MID)
    
    y = PAGE_H - HEADER_H - 8 * mm
    
    # ── PRODUCT TITLE ────────────────────────────────────────────────────────
    prod_name = offer_data.get('product_name', '')
    manufacturer = offer_data.get('manufacturer', '')
    cat_text = offer_data.get('product_category', 'Sprzęt medyczny')
    cat_key = offer_data.get('cat_key', 'USG')
    
    title_size = 20
    c.saveState()
    c.setFont('Inter-Bold', title_size)
    tw = c.stringWidth(prod_name, 'Inter-Bold', title_size)
    while tw > CONTENT_W and title_size > 14:
        title_size -= 1
        tw = c.stringWidth(prod_name, 'Inter-Bold', title_size)
    c.restoreState()
    
    draw_text(c, prod_name, MARGIN_L, y, font='Inter-Bold', size=title_size, color=C_DARK)
    
    # Category + manufacturer badges (below title)
    badge_y = y - title_size - 4
    bx = MARGIN_L
    badge_w = pill_badge(c, bx, badge_y - 10, cat_text, accent, white, 'Inter-SemiBold', 7.5)
    
    if manufacturer:
        mx = bx + badge_w + 6
        c.saveState()
        c.setFillColor(C_PANEL)
        c.setStrokeColor(C_LINE)
        c.setLineWidth(0.5)
        mw = c.stringWidth(manufacturer, 'Inter', 7.5) + 10
        c.roundRect(mx, badge_y - 10, mw, 13, 6, fill=1, stroke=1)
        c.setFillColor(C_DARK)
        c.setFont('Inter', 7.5)
        c.drawString(mx + 5, badge_y - 6.5, manufacturer)
        c.restoreState()
    
    y -= title_size + 22
    
    # ── IMAGE + SPEC TABLE ───────────────────────────────────────────────────
    img_col_w = CONTENT_W * 0.42
    spec_col_w = CONTENT_W * 0.56
    spec_col_x = MARGIN_L + img_col_w + CONTENT_W * 0.02
    
    # Spec table height estimate
    specs = offer_data.get('specs', [])
    row_h = 11.5
    spec_table_h = len(specs) * row_h + 4 + 13  # +13 for header
    
    img_h = max(spec_table_h, 68 * mm)  # at least 68mm tall
    img_y = y - img_h
    
    # Image panel
    prod_img = offer_data.get('product_image', '')
    img_path = prod_img if prod_img else os.path.join(IMG_DIR, f"{offer_data.get('cat_key_img','')}.png")
    draw_product_image(c, img_path, MARGIN_L, img_y, img_col_w, img_h)
    
    # Spec table
    ty = y
    # Header row
    draw_rect(c, spec_col_x, ty - 13, spec_col_w, 13, fill=C_DARK)
    draw_text(c, 'SPECYFIKACJA TECHNICZNA', spec_col_x + 5, ty - 9.5,
              font='Inter-SemiBold', size=7.5, color=white)
    ty -= 13
    
    for i, (param, value) in enumerate(specs):
        row_bg = C_PANEL if i % 2 == 0 else white
        draw_rect(c, spec_col_x, ty - row_h, spec_col_w, row_h, fill=row_bg)
        
        # Left bar for alt rows
        if i % 2 == 0:
            draw_rect(c, spec_col_x, ty - row_h, 2.5, row_h, fill=HexColor('#DDE8F0'))
        
        # Param name
        param_w = spec_col_w * 0.36
        draw_text(c, str(param), spec_col_x + 6, ty - row_h + 3,
                  font='Inter-SemiBold', size=7, color=C_DARK)
        
        # Value (with truncation if needed)
        val_x = spec_col_x + param_w
        val_w = spec_col_w - param_w - 4
        c.saveState()
        c.setFont('Inter', 7)
        val_str = str(value)
        while c.stringWidth(val_str, 'Inter', 7) > val_w and len(val_str) > 4:
            val_str = val_str[:-2] + '…'
        c.setFillColor(C_TEXT)
        c.drawString(val_x, ty - row_h + 3, val_str)
        c.restoreState()
        
        ty -= row_h
    
    # Bottom spec border
    draw_line(c, spec_col_x, ty, spec_col_x + spec_col_w, ty, color=C_LINE)
    
    y = min(img_y, ty) - 8 * mm
    
    # ── DESCRIPTION + BENEFITS ───────────────────────────────────────────────
    desc_w = CONTENT_W * 0.55
    ben_w = CONTENT_W * 0.42
    ben_x = MARGIN_L + desc_w + CONTENT_W * 0.03
    
    # Description section
    section_y = y
    section_title(c, MARGIN_L, section_y, 'OPIS PRODUKTU', color=accent, width=desc_w)
    section_y -= 6
    
    desc_text = offer_data.get('description', '')
    y_after_desc = draw_wrapped_text(c, desc_text, MARGIN_L, section_y, desc_w,
                                     font='Inter', size=8.5, color=C_TEXT, line_height=13)
    
    # Benefits section
    ben_y = y
    section_title(c, ben_x, ben_y, 'KLUCZOWE KORZYŚCI', color=accent, width=ben_w)
    ben_y -= 6
    
    benefits = offer_data.get('benefits', [])
    for benefit in benefits:
        if ben_y < MARGIN_B + FOOTER_H + 10:
            break
        # Check mark circle
        c.saveState()
        c.setFillColor(accent)
        c.circle(ben_x + 4.5, ben_y - 2, 3.5, fill=1, stroke=0)
        c.setFont('Inter-Bold', 5.5)
        c.setFillColor(white)
        c.drawCentredString(ben_x + 4.5, ben_y - 4, '✓')
        c.restoreState()
        
        ben_y = draw_wrapped_text(c, benefit, ben_x + 12, ben_y, ben_w - 12,
                                  font='Inter', size=8.5, color=C_TEXT, line_height=12)
        ben_y -= 2
    
    c.showPage()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 3 — DETAILED SPECS / ACCESSORIES / FEATURES
# ══════════════════════════════════════════════════════════════════════════════

def build_details_page(c, offer_data):
    draw_header(c, offer_data)
    draw_footer(c, page_num=3, total_pages=4)
    
    accent = CAT_COLORS.get(offer_data.get('cat_key', 'USG'), C_MID)
    y = PAGE_H - HEADER_H - 10 * mm
    
    # ── DETAILED SPECS ───────────────────────────────────────────────────────
    detailed_specs = offer_data.get('detailed_specs') or offer_data.get('specs', [])
    
    section_title(c, MARGIN_L, y, 'SZCZEGÓŁOWA SPECYFIKACJA TECHNICZNA', color=accent, width=CONTENT_W)
    y -= 8
    
    row_h = 11
    col1_w = CONTENT_W * 0.35
    col2_w = CONTENT_W - col1_w
    
    # Table header
    draw_rect(c, MARGIN_L, y - 12, CONTENT_W, 12, fill=C_DARK)
    draw_text(c, 'Parametr', MARGIN_L + 6, y - 8.5, font='Inter-SemiBold', size=7.5, color=white)
    draw_text(c, 'Wartość', MARGIN_L + col1_w + 6, y - 8.5, font='Inter-SemiBold', size=7.5, color=white)
    y -= 12
    
    cutoff = MARGIN_B + FOOTER_H + 40 * mm  # leave space for accessories section
    
    for i, (param, value) in enumerate(detailed_specs):
        if y - row_h < cutoff:
            break
        row_bg = C_PANEL if i % 2 == 0 else white
        draw_rect(c, MARGIN_L, y - row_h, CONTENT_W, row_h, fill=row_bg)
        if i % 2 == 0:
            draw_rect(c, MARGIN_L, y - row_h, 2.5, row_h, fill=HexColor('#DDE8F0'))
        
        draw_text(c, str(param), MARGIN_L + 6, y - row_h + 3,
                  font='Inter-SemiBold', size=7.5, color=C_DARK)
        
        c.saveState()
        c.setFont('Inter', 7.5)
        val_str = str(value)
        max_val_w = col2_w - 10
        while c.stringWidth(val_str, 'Inter', 7.5) > max_val_w and len(val_str) > 4:
            val_str = val_str[:-2] + '…'
        c.setFillColor(C_TEXT)
        c.drawString(MARGIN_L + col1_w + 6, y - row_h + 3, val_str)
        c.restoreState()
        
        y -= row_h
    
    draw_line(c, MARGIN_L, y, MARGIN_L + CONTENT_W, y, color=C_LINE)
    y -= 8 * mm
    
    # ── TWO COLUMNS: ACCESSORIES | EXTRA FEATURES ───────────────────────────
    half_w = CONTENT_W * 0.5 - 4 * mm
    right_x = MARGIN_L + half_w + 8 * mm
    
    accessories = offer_data.get('accessories', [])
    extra_features = offer_data.get('extra_features', [])
    
    col_y = y
    left_col_y = col_y
    right_col_y = col_y
    
    # LEFT: Accessories
    if accessories:
        section_title(c, MARGIN_L, left_col_y, 'DOSTĘPNE AKCESORIA I KONFIGURACJE', color=accent, width=half_w)
        left_col_y -= 8
        
        # Sub-header
        draw_rect(c, MARGIN_L, left_col_y - 11, half_w, 11, fill=C_MID)
        draw_text(c, 'Akcesoria', MARGIN_L + 6, left_col_y - 7.5, font='Inter-SemiBold', size=7, color=white)
        draw_text(c, 'Cena / Info', MARGIN_L + half_w - 50, left_col_y - 7.5, font='Inter-SemiBold', size=7, color=white)
        left_col_y -= 11
        
        for i, (name, price) in enumerate(accessories):
            row_bg = C_PANEL if i % 2 == 0 else white
            draw_rect(c, MARGIN_L, left_col_y - row_h, half_w, row_h, fill=row_bg)
            
            c.saveState()
            c.setFont('Inter', 7.5)
            aname = str(name)
            max_aw = half_w - 65
            while c.stringWidth(aname, 'Inter', 7.5) > max_aw and len(aname) > 4:
                aname = aname[:-2] + '…'
            c.setFillColor(C_TEXT)
            c.drawString(MARGIN_L + 5, left_col_y - row_h + 3, aname)
            
            # Price right-aligned
            c.setFont('Inter-SemiBold', 7.5)
            c.setFillColor(C_MID)
            c.drawRightString(MARGIN_L + half_w - 4, left_col_y - row_h + 3, str(price))
            c.restoreState()
            
            left_col_y -= row_h
        
        draw_line(c, MARGIN_L, left_col_y, MARGIN_L + half_w, left_col_y, color=C_LINE)
    
    # RIGHT: Extra features
    if extra_features:
        section_title(c, right_x, right_col_y, 'CECHY WYRÓŻNIAJĄCE', color=accent, width=half_w)
        right_col_y -= 8
        
        for feat in extra_features:
            if right_col_y < MARGIN_B + FOOTER_H + 5:
                break
            # Bullet
            draw_rect(c, right_x, right_col_y - 6, 3.5, 3.5, fill=accent, radius=1)
            right_col_y = draw_wrapped_text(c, str(feat), right_x + 8, right_col_y, 
                                             half_w - 8, font='Inter', size=8.5, 
                                             color=C_TEXT, line_height=12.5)
            right_col_y -= 3
    
    c.showPage()

# ══════════════════════════════════════════════════════════════════════════════
# PAGE 4 — PRICING
# ══════════════════════════════════════════════════════════════════════════════

def build_pricing_page(c, offer_data):
    draw_header(c, offer_data)
    draw_footer(c, page_num=4, total_pages=4)
    
    accent = CAT_COLORS.get(offer_data.get('cat_key', 'USG'), C_MID)
    y = PAGE_H - HEADER_H - 10 * mm
    
    # ── TITLE ────────────────────────────────────────────────────────────────
    section_title(c, MARGIN_L, y, 'OFERTA CENOWA', color=accent, width=CONTENT_W)
    y -= 14
    
    # Client info
    client = offer_data.get('client_name', '')
    client_addr = offer_data.get('client_address', '')
    if client:
        draw_text(c, f'Przygotowana dla: {client}', MARGIN_L, y,
                  font='Inter-SemiBold', size=8.5, color=C_DARK)
        if client_addr:
            y -= 12
            draw_text(c, client_addr, MARGIN_L, y, font='Inter', size=8, color=C_CAPTION)
        y -= 14
    
    # ── PRICING TABLE ────────────────────────────────────────────────────────
    pricing = offer_data.get('pricing', [])
    vat_rate = offer_data.get('vat_rate', 0.08)
    
    # Column widths
    cols = {
        'model': CONTENT_W * 0.38,
        'desc':  CONTENT_W * 0.22,
        'qty':   CONTENT_W * 0.07,
        'unit':  CONTENT_W * 0.16,
        'total': CONTENT_W * 0.17,
    }
    
    # Header
    header_h = 13
    draw_rect(c, MARGIN_L, y - header_h, CONTENT_W, header_h, fill=C_DARK)
    cx = MARGIN_L
    headers = [('Model / Konfiguracja', cols['model']),
               ('Opis', cols['desc']),
               ('Ilość', cols['qty']),
               ('Cena netto (PLN)', cols['unit']),
               ('Wartość netto (PLN)', cols['total'])]
    for htext, hw in headers:
        draw_text(c, htext, cx + 4, y - header_h + 4, font='Inter-SemiBold', size=7, color=white)
        cx += hw
    y -= header_h
    
    grand_total_netto = 0
    row_h = 14
    
    for i, item in enumerate(pricing):
        model = item.get('model', '')
        desc = item.get('description', '')
        qty = item.get('qty', 1)
        unit_netto = item.get('unit_price_netto', 0)
        total_netto = qty * unit_netto
        grand_total_netto += total_netto
        
        row_bg = C_PANEL if i % 2 == 0 else white
        draw_rect(c, MARGIN_L, y - row_h, CONTENT_W, row_h, fill=row_bg)
        
        # Left accent bar
        draw_rect(c, MARGIN_L, y - row_h, 3, row_h, fill=accent)
        
        cx = MARGIN_L
        # Model
        c.saveState()
        c.setFont('Inter-SemiBold', 8)
        ms = model
        while c.stringWidth(ms, 'Inter-SemiBold', 8) > cols['model'] - 10 and len(ms) > 4:
            ms = ms[:-2] + '…'
        c.setFillColor(C_DARK)
        c.drawString(cx + 6, y - row_h + 4, ms)
        c.restoreState()
        cx += cols['model']
        
        # Description
        c.saveState()
        c.setFont('Inter', 7.5)
        ds = desc
        while c.stringWidth(ds, 'Inter', 7.5) > cols['desc'] - 8 and len(ds) > 4:
            ds = ds[:-2] + '…'
        c.setFillColor(C_TEXT)
        c.drawString(cx + 4, y - row_h + 4, ds)
        c.restoreState()
        cx += cols['desc']
        
        # Qty
        draw_text(c, str(qty), cx + cols['qty']/2, y - row_h + 4, 
                  font='Inter', size=8, color=C_TEXT, align='center')
        cx += cols['qty']
        
        # Unit price
        draw_text(c, format_price(unit_netto), cx + cols['unit'] - 4, y - row_h + 4,
                  font='Inter-SemiBold', size=8, color=C_DARK, align='right')
        cx += cols['unit']
        
        # Total
        draw_text(c, format_price(total_netto), cx + cols['total'] - 4, y - row_h + 4,
                  font='Inter-Bold', size=8, color=C_MID, align='right')
        
        y -= row_h
    
    # ── TOTALS BOX ───────────────────────────────────────────────────────────
    draw_line(c, MARGIN_L, y, MARGIN_L + CONTENT_W, y, color=C_LINE, width=1)
    y -= 5
    
    vat_amount = grand_total_netto * vat_rate
    total_brutto = grand_total_netto + vat_amount
    vat_pct = int(vat_rate * 100)
    
    totals_w = CONTENT_W * 0.45
    totals_x = MARGIN_L + CONTENT_W - totals_w
    
    # Summary rows
    summary_rows = [
        ('Razem netto:', format_price(grand_total_netto) + ' PLN', False),
        (f'VAT {vat_pct}%:', format_price(vat_amount) + ' PLN', False),
        ('RAZEM BRUTTO:', format_price(total_brutto) + ' PLN', True),
    ]
    
    for label, amount, is_total in summary_rows:
        row_h2 = 13 if not is_total else 16
        if is_total:
            draw_rect(c, totals_x, y - row_h2, totals_w, row_h2, fill=C_DARK)
            draw_text(c, label, totals_x + 8, y - row_h2 + 4.5,
                      font='Inter-Bold', size=9, color=white)
            draw_text(c, amount, totals_x + totals_w - 8, y - row_h2 + 4.5,
                      font='Inter-Bold', size=10, color=white, align='right')
        else:
            draw_rect(c, totals_x, y - row_h2, totals_w, row_h2, fill=C_PANEL)
            draw_line(c, totals_x, y, totals_x + totals_w, y, color=C_LINE)
            draw_text(c, label, totals_x + 8, y - row_h2 + 4,
                      font='Inter', size=8.5, color=C_CAPTION)
            draw_text(c, amount, totals_x + totals_w - 8, y - row_h2 + 4,
                      font='Inter-SemiBold', size=8.5, color=C_DARK, align='right')
        y -= row_h2
    
    y -= 8 * mm
    
    # ── TERMS SECTION ────────────────────────────────────────────────────────
    terms_w = CONTENT_W * 0.55
    
    section_title(c, MARGIN_L, y, 'WARUNKI OFERTY', color=accent, width=terms_w)
    y -= 20
    
    terms = [
        ('Warunki płatności:', offer_data.get('payment_terms', '—')),
        ('Czas dostawy:', offer_data.get('delivery_time', '—')),
        ('Gwarancja:', offer_data.get('warranty', '—')),
        ('Oferta ważna do:', offer_data.get('valid_until', '—')),
    ]
    
    for label, value in terms:
        draw_text(c, label, MARGIN_L, y, font='Inter-SemiBold', size=8.5, color=C_DARK)
        draw_text(c, value, MARGIN_L + 55*mm, y, font='Inter', size=8.5, color=C_TEXT)
        y -= 13
    
    y -= 4
    
    # Notes
    notes = offer_data.get('notes', [])
    if notes:
        for note in notes:
            draw_text(c, f'* {note}', MARGIN_L, y, font='Inter-Light', size=7.5, color=C_CAPTION)
            y -= 11
    
    y -= 8 * mm
    
    # ── SIGNATURE SECTION ────────────────────────────────────────────────────
    sig_y = MARGIN_B + FOOTER_H + 38 * mm
    
    if y > sig_y + 5:
        # Use y position
        sig_y = y
    
    sig_col_w = CONTENT_W * 0.42
    sig_r_x = MARGIN_L + CONTENT_W - sig_col_w
    
    # Left: company stamp area
    draw_rect(c, MARGIN_L, sig_y - 30*mm, sig_col_w, 28*mm, 
              fill=None, stroke=C_LINE, stroke_w=0.5, radius=3)
    draw_text(c, 'Pieczątka i podpis LaudaMed', MARGIN_L + sig_col_w/2, sig_y - 5*mm,
              font='Inter-Light', size=7.5, color=C_CAPTION, align='center')
    draw_line(c, MARGIN_L + 8, sig_y - 25*mm, MARGIN_L + sig_col_w - 8, sig_y - 25*mm,
              color=C_LINE)
    draw_text(c, 'Data:', MARGIN_L + 8, sig_y - 28*mm + 3,
              font='Inter', size=7.5, color=C_CAPTION)
    
    # Right: client signature area
    draw_rect(c, sig_r_x, sig_y - 30*mm, sig_col_w, 28*mm,
              fill=None, stroke=C_LINE, stroke_w=0.5, radius=3)
    draw_text(c, 'Pieczątka i podpis Zamawiającego', sig_r_x + sig_col_w/2, sig_y - 5*mm,
              font='Inter-Light', size=7.5, color=C_CAPTION, align='center')
    draw_line(c, sig_r_x + 8, sig_y - 25*mm, sig_r_x + sig_col_w - 8, sig_y - 25*mm,
              color=C_LINE)
    draw_text(c, 'Data:', sig_r_x + 8, sig_y - 28*mm + 3,
              font='Inter', size=7.5, color=C_CAPTION)
    
    c.showPage()

# ══════════════════════════════════════════════════════════════════════════════
# PDF GENERATOR — MAIN ENTRY
# ══════════════════════════════════════════════════════════════════════════════

def generate_pdf(offer_data: dict, output_path: str):
    """Generate a 4-page PDF sales offer."""
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    register_fonts()
    
    c = canvas.Canvas(output_path, pagesize=A4)
    c.setAuthor(COMPANY['name'])
    c.setTitle(f"Oferta Handlowa — {offer_data.get('product_name', '')}")
    c.setSubject(offer_data.get('offer_number', ''))
    
    print(f"  Page 1: Cover...")
    build_cover_page(c, offer_data)
    
    print(f"  Page 2: Product card...")
    build_product_card_page(c, offer_data)
    
    print(f"  Page 3: Details / Accessories...")
    build_details_page(c, offer_data)
    
    print(f"  Page 4: Pricing...")
    build_pricing_page(c, offer_data)
    
    c.save()
    print(f"[OK] PDF saved: {output_path}")
    return output_path

# ══════════════════════════════════════════════════════════════════════════════
# DOCX TEMPLATE GENERATOR
# ══════════════════════════════════════════════════════════════════════════════

def _set_doc_margins(doc, top=18, bottom=18, left=18, right=18):
    """Set document margins in mm."""
    from docx.oxml import OxmlElement
    sections = doc.sections
    for section in sections:
        section.top_margin = Mm(top)
        section.bottom_margin = Mm(bottom)
        section.left_margin = Mm(left)
        section.right_margin = Mm(right)
        section.page_width = Mm(210)
        section.page_height = Mm(297)

def _set_cell_bg(cell, hex_color: str):
    """Set table cell background color."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    shd = OxmlElement('w:shd')
    shd.set(qn('w:val'), 'clear')
    shd.set(qn('w:color'), 'auto')
    shd.set(qn('w:fill'), hex_color.lstrip('#'))
    tcPr.append(shd)

def _set_cell_border(cell, **kwargs):
    """Set cell borders."""
    tc = cell._tc
    tcPr = tc.get_or_add_tcPr()
    tcBorders = OxmlElement('w:tcBorders')
    for edge in ('top', 'bottom', 'left', 'right', 'insideH', 'insideV'):
        if edge in kwargs:
            tag = OxmlElement(f'w:{edge}')
            tag.set(qn('w:val'), kwargs[edge].get('val', 'single'))
            tag.set(qn('w:sz'), str(kwargs[edge].get('sz', 4)))
            tag.set(qn('w:color'), kwargs[edge].get('color', '000000'))
            tcBorders.append(tag)
    tcPr.append(tcBorders)

def _para_style(para, font_name='Calibri', size_pt=10, bold=False, 
                color_hex=None, align=WD_ALIGN_PARAGRAPH.LEFT, space_before=0, space_after=4):
    para.alignment = align
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(space_after)
    for run in para.runs:
        run.font.name = font_name
        run.font.size = Pt(size_pt)
        run.font.bold = bold
        if color_hex:
            r, g, b = bytes.fromhex(color_hex.lstrip('#'))
            run.font.color.rgb = RGBColor(r, g, b)

def add_styled_heading(doc, text, level=1, color='0D2137', size=14, bold=True, space_before=10):
    para = doc.add_paragraph()
    para.paragraph_format.space_before = Pt(space_before)
    para.paragraph_format.space_after = Pt(4)
    run = para.add_run(text)
    run.font.name = 'Calibri'
    run.font.size = Pt(size)
    run.font.bold = bold
    r, g, b = bytes.fromhex(color.lstrip('#'))
    run.font.color.rgb = RGBColor(r, g, b)
    return para

def add_two_col_table(doc, rows, col1_w_pct=0.35, header_row=None, 
                      header_bg='0D2137', alt_bg='F5F8FA'):
    """Add a two-column table (param / value)."""
    table = doc.add_table(rows=0, cols=2)
    table.alignment = WD_TABLE_ALIGNMENT.LEFT
    table.style = 'Table Grid'
    
    total_w = Mm(210 - 36)  # page width minus margins
    col1_w = int(total_w * col1_w_pct)
    col2_w = int(total_w * (1 - col1_w_pct))
    
    if header_row:
        row = table.add_row()
        row.cells[0].merge(row.cells[1])
        row.cells[0].text = header_row
        _set_cell_bg(row.cells[0], header_bg)
        p = row.cells[0].paragraphs[0]
        run = p.runs[0] if p.runs else p.add_run(header_row)
        run.font.color.rgb = RGBColor(255, 255, 255)
        run.font.bold = True
        run.font.size = Pt(9)
        run.font.name = 'Calibri'
    
    for i, (param, value) in enumerate(rows):
        row = table.add_row()
        row.cells[0].text = str(param)
        row.cells[1].text = str(value)
        
        bg = alt_bg if i % 2 == 0 else 'FFFFFF'
        _set_cell_bg(row.cells[0], bg)
        _set_cell_bg(row.cells[1], bg)
        
        for cell in row.cells:
            for para in cell.paragraphs:
                para.paragraph_format.space_before = Pt(1)
                para.paragraph_format.space_after = Pt(1)
                for run in para.runs:
                    run.font.size = Pt(9)
                    run.font.name = 'Calibri'
        
        # Bold param name
        for run in row.cells[0].paragraphs[0].runs:
            run.font.bold = True
            run.font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
    
    # Set column widths
    for row in table.rows:
        for i, cell in enumerate(row.cells):
            cell.width = col1_w if i == 0 else col2_w
    
    return table

def generate_docx_template(output_path: str):
    """Generate a DOCX template with placeholder fields."""
    if not DOCX_AVAILABLE:
        print("[error] python-docx not installed. Run: pip install python-docx")
        return None
    
    os.makedirs(os.path.dirname(output_path), exist_ok=True)
    doc = Document()
    _set_doc_margins(doc, 18, 18, 20, 20)
    
    # ── Style defaults ────────────────────────────────────────────────────────
    style = doc.styles['Normal']
    style.font.name = 'Calibri'
    style.font.size = Pt(10)
    
    # ════════════════════════════════════════════════════════════════
    # PAGE 1 — COVER (simulated as a styled section)
    # ════════════════════════════════════════════════════════════════
    
    # Title block with dark background (simulated via table)
    cover_table = doc.add_table(rows=1, cols=1)
    cover_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    cell = cover_table.rows[0].cells[0]
    cell.width = Mm(174)
    _set_cell_bg(cell, '0D2137')
    
    p = cell.add_paragraph()
    p.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p.paragraph_format.space_before = Pt(18)
    p.paragraph_format.space_after = Pt(4)
    run = p.add_run('LaudaMed Sp. z o.o.')
    run.font.name = 'Calibri'
    run.font.size = Pt(11)
    run.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    
    p2 = cell.add_paragraph()
    p2.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p2.paragraph_format.space_before = Pt(4)
    p2.paragraph_format.space_after = Pt(4)
    run2 = p2.add_run('OFERTA HANDLOWA')
    run2.font.name = 'Calibri'
    run2.font.size = Pt(10)
    run2.font.color.rgb = RGBColor(0xCC, 0xDD, 0xFF)
    run2.font.bold = True
    
    p3 = cell.add_paragraph()
    p3.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p3.paragraph_format.space_before = Pt(8)
    p3.paragraph_format.space_after = Pt(4)
    run3 = p3.add_run('[PRODUCT_NAME]')
    run3.font.name = 'Calibri'
    run3.font.size = Pt(24)
    run3.font.bold = True
    run3.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    p4 = cell.add_paragraph()
    p4.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p4.paragraph_format.space_before = Pt(4)
    p4.paragraph_format.space_after = Pt(4)
    run4 = p4.add_run('[PRODUCT_CATEGORY]  ·  [MANUFACTURER]')
    run4.font.name = 'Calibri'
    run4.font.size = Pt(11)
    run4.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    
    # Offer details row
    info_table = doc.add_table(rows=1, cols=3)
    info_table.alignment = WD_TABLE_ALIGNMENT.CENTER
    _set_cell_bg(info_table.rows[0].cells[0], '1A3050')
    _set_cell_bg(info_table.rows[0].cells[1], '1A3050')
    _set_cell_bg(info_table.rows[0].cells[2], '1A3050')
    
    info_labels = [('Nr oferty:', '[OFFER_NUMBER]'), 
                   ('Data:', '[DATE]'), 
                   ('Ważna do:', '[VALID_UNTIL]')]
    for j, (lbl, val) in enumerate(info_labels):
        c_cell = info_table.rows[0].cells[j]
        pl = c_cell.add_paragraph()
        pl.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pl.paragraph_format.space_before = Pt(3)
        pl.paragraph_format.space_after = Pt(1)
        rl = pl.add_run(lbl)
        rl.font.size = Pt(7)
        rl.font.name = 'Calibri'
        rl.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
        
        pv = c_cell.add_paragraph()
        pv.alignment = WD_ALIGN_PARAGRAPH.CENTER
        pv.paragraph_format.space_before = Pt(1)
        pv.paragraph_format.space_after = Pt(4)
        rv = pv.add_run(val)
        rv.font.size = Pt(9)
        rv.font.bold = True
        rv.font.name = 'Calibri'
        rv.font.color.rgb = RGBColor(0xFF, 0xFF, 0xFF)
    
    doc.add_paragraph()
    
    # Client info
    add_styled_heading(doc, 'Dane Klienta', level=2, color='1565A0', size=11, space_before=4)
    client_table = doc.add_table(rows=2, cols=2)
    client_table.style = 'Table Grid'
    labels_vals = [('Nazwa klienta:', '[CLIENT_NAME]'), ('Adres:', '[CLIENT_ADDRESS]')]
    for i, (lbl, val) in enumerate(labels_vals):
        client_table.rows[i].cells[0].text = lbl
        client_table.rows[i].cells[1].text = val
        for cell in client_table.rows[i].cells:
            for p in cell.paragraphs:
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.name = 'Calibri'
        client_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
    
    # Page break
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # PAGE 2 — PRODUCT CARD
    # ════════════════════════════════════════════════════════════════
    
    # Header simulation
    hdr_table = doc.add_table(rows=1, cols=2)
    hdr_table.rows[0].height = Mm(12)
    _set_cell_bg(hdr_table.rows[0].cells[0], '0D2137')
    _set_cell_bg(hdr_table.rows[0].cells[1], '0D2137')
    
    p_logo = hdr_table.rows[0].cells[0].add_paragraph()
    p_logo.paragraph_format.space_before = Pt(3)
    p_logo.paragraph_format.space_after = Pt(2)
    rl = p_logo.add_run('LaudaMed')
    rl.font.bold = True
    rl.font.name = 'Calibri'
    rl.font.size = Pt(11)
    rl.font.color.rgb = RGBColor(255, 255, 255)
    
    p_hdr_r = hdr_table.rows[0].cells[1].add_paragraph()
    p_hdr_r.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    p_hdr_r.paragraph_format.space_before = Pt(3)
    rr = p_hdr_r.add_run('OFERTA HANDLOWA | [YEAR]  ·  [OFFER_NUMBER]')
    rr.font.size = Pt(8)
    rr.font.name = 'Calibri'
    rr.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    
    doc.add_paragraph()
    
    # Product title
    add_styled_heading(doc, '[PRODUCT_NAME]', level=1, color='0D2137', size=20, space_before=2)
    
    p_sub = doc.add_paragraph()
    p_sub.paragraph_format.space_before = Pt(0)
    p_sub.paragraph_format.space_after = Pt(6)
    rs = p_sub.add_run('[PRODUCT_CATEGORY]  ·  [MANUFACTURER]  ·  Model: [MODEL]')
    rs.font.size = Pt(9)
    rs.font.name = 'Calibri'
    rs.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Image placeholder
    img_table = doc.add_table(rows=1, cols=2)
    img_table.style = 'Table Grid'
    img_cell = img_table.rows[0].cells[0]
    img_cell.width = Mm(80)
    _set_cell_bg(img_cell, 'F5F8FA')
    img_para = img_cell.add_paragraph()
    img_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
    img_para.paragraph_format.space_before = Pt(30)
    img_para.paragraph_format.space_after = Pt(30)
    img_run = img_para.add_run('[PRODUCT_IMAGE]')
    img_run.font.size = Pt(9)
    img_run.font.name = 'Calibri'
    img_run.font.color.rgb = RGBColor(0x99, 0x99, 0x99)
    
    spec_cell = img_table.rows[0].cells[1]
    spec_cell.width = Mm(94)
    spec_para = spec_cell.add_paragraph()
    spec_para.paragraph_format.space_before = Pt(2)
    spec_run = spec_para.add_run('[SPECS_TABLE — wstaw tabelę specyfikacji]')
    spec_run.font.size = Pt(9)
    spec_run.font.italic = True
    spec_run.font.name = 'Calibri'
    spec_run.font.color.rgb = RGBColor(0x66, 0x88, 0xAA)
    
    doc.add_paragraph()
    
    # Description
    add_styled_heading(doc, 'OPIS PRODUKTU', level=2, color='1565A0', size=11, space_before=6)
    p_desc = doc.add_paragraph('[DESCRIPTION]')
    p_desc.paragraph_format.space_before = Pt(2)
    p_desc.paragraph_format.space_after = Pt(6)
    p_desc.runs[0].font.size = Pt(9.5)
    p_desc.runs[0].font.name = 'Calibri'
    p_desc.runs[0].font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Benefits
    add_styled_heading(doc, 'KLUCZOWE KORZYŚCI', level=2, color='1565A0', size=11, space_before=6)
    for i in range(1, 7):
        p_ben = doc.add_paragraph(f'[BENEFIT_{i}]', style='List Bullet')
        p_ben.paragraph_format.space_before = Pt(1)
        p_ben.paragraph_format.space_after = Pt(1)
        if p_ben.runs:
            p_ben.runs[0].font.size = Pt(9.5)
            p_ben.runs[0].font.name = 'Calibri'
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # PAGE 3 — DETAILED SPECS / ACCESSORIES
    # ════════════════════════════════════════════════════════════════
    
    # Header
    hdr2 = doc.add_table(rows=1, cols=2)
    _set_cell_bg(hdr2.rows[0].cells[0], '0D2137')
    _set_cell_bg(hdr2.rows[0].cells[1], '0D2137')
    pl2 = hdr2.rows[0].cells[0].add_paragraph()
    pl2.paragraph_format.space_before = Pt(3)
    rl2 = pl2.add_run('LaudaMed')
    rl2.font.bold = True; rl2.font.name = 'Calibri'; rl2.font.size = Pt(11)
    rl2.font.color.rgb = RGBColor(255, 255, 255)
    pr2 = hdr2.rows[0].cells[1].add_paragraph()
    pr2.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr2 = pr2.add_run('OFERTA HANDLOWA | [OFFER_NUMBER]')
    rr2.font.size = Pt(8); rr2.font.name = 'Calibri'
    rr2.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    
    doc.add_paragraph()
    
    add_styled_heading(doc, 'SZCZEGÓŁOWA SPECYFIKACJA TECHNICZNA', level=2, 
                       color='0D2137', size=13, space_before=2)
    
    placeholder_specs = [
        ('Klasa', '[SPEC_VALUE]'),
        ('Platforma', '[SPEC_VALUE]'),
        ('Monitor', '[SPEC_VALUE]'),
        ('Głowice', '[SPEC_VALUE]'),
        ('Tryby obrazowania', '[SPEC_VALUE]'),
        ('Technologie', '[SPEC_VALUE]'),
        ('Interfejsy', '[SPEC_VALUE]'),
        ('Wymiary', '[SPEC_VALUE]'),
        ('Gwarancja', '[SPEC_VALUE]'),
    ]
    add_two_col_table(doc, placeholder_specs, header_row='SPECYFIKACJA TECHNICZNA — [PRODUCT_NAME]')
    
    doc.add_paragraph()
    add_styled_heading(doc, 'DOSTĘPNE AKCESORIA I KONFIGURACJE', level=2,
                       color='0D2137', size=13, space_before=4)
    
    placeholder_acc = [
        ('[ACCESSORY_1_NAME]', '[PRICE]'),
        ('[ACCESSORY_2_NAME]', '[PRICE]'),
        ('[ACCESSORY_3_NAME]', '[PRICE]'),
    ]
    add_two_col_table(doc, placeholder_acc, col1_w_pct=0.6, 
                      header_row='Akcesorium / Konfiguracja  |  Cena netto')
    
    doc.add_paragraph()
    add_styled_heading(doc, 'CECHY WYRÓŻNIAJĄCE', level=2, color='0D2137', size=13, space_before=4)
    for i in range(1, 5):
        p_feat = doc.add_paragraph(f'[FEATURE_{i}]', style='List Bullet')
        p_feat.paragraph_format.space_before = Pt(1)
        p_feat.paragraph_format.space_after = Pt(1)
        if p_feat.runs:
            p_feat.runs[0].font.size = Pt(9.5)
            p_feat.runs[0].font.name = 'Calibri'
    
    doc.add_page_break()
    
    # ════════════════════════════════════════════════════════════════
    # PAGE 4 — PRICING
    # ════════════════════════════════════════════════════════════════
    
    # Header
    hdr3 = doc.add_table(rows=1, cols=2)
    _set_cell_bg(hdr3.rows[0].cells[0], '0D2137')
    _set_cell_bg(hdr3.rows[0].cells[1], '0D2137')
    pl3 = hdr3.rows[0].cells[0].add_paragraph()
    pl3.paragraph_format.space_before = Pt(3)
    rl3 = pl3.add_run('LaudaMed')
    rl3.font.bold = True; rl3.font.name = 'Calibri'; rl3.font.size = Pt(11)
    rl3.font.color.rgb = RGBColor(255, 255, 255)
    pr3 = hdr3.rows[0].cells[1].add_paragraph()
    pr3.alignment = WD_ALIGN_PARAGRAPH.RIGHT
    rr3 = pr3.add_run('OFERTA HANDLOWA | [OFFER_NUMBER]')
    rr3.font.size = Pt(8); rr3.font.name = 'Calibri'
    rr3.font.color.rgb = RGBColor(0xAA, 0xCC, 0xFF)
    
    doc.add_paragraph()
    add_styled_heading(doc, 'OFERTA CENOWA', level=1, color='0D2137', size=16, space_before=2)
    
    p_cli = doc.add_paragraph()
    p_cli.paragraph_format.space_before = Pt(0)
    p_cli.paragraph_format.space_after = Pt(8)
    rc = p_cli.add_run('Przygotowana dla: [CLIENT_NAME]')
    rc.font.size = Pt(10); rc.font.name = 'Calibri'
    rc.font.color.rgb = RGBColor(0x15, 0x65, 0xA0)
    
    # Pricing table
    price_table = doc.add_table(rows=1, cols=5)
    price_table.style = 'Table Grid'
    price_table.alignment = WD_TABLE_ALIGNMENT.LEFT
    
    price_headers = ['Model / Konfiguracja', 'Opis', 'Ilość', 'Cena netto (PLN)', 'Wartość netto (PLN)']
    for j, ph in enumerate(price_headers):
        _set_cell_bg(price_table.rows[0].cells[j], '0D2137')
        p_ph = price_table.rows[0].cells[j].add_paragraph()
        r_ph = p_ph.add_run(ph)
        r_ph.font.size = Pt(8); r_ph.font.bold = True; r_ph.font.name = 'Calibri'
        r_ph.font.color.rgb = RGBColor(255, 255, 255)
        p_ph.paragraph_format.space_before = Pt(2)
        p_ph.paragraph_format.space_after = Pt(2)
    
    # One data row placeholder
    data_row = price_table.add_row()
    placeholders = ['[MODEL_NAME]', '[CONFIGURATION]', '[QTY]', '[UNIT_PRICE]', '[TOTAL_PRICE]']
    for j, ph in enumerate(placeholders):
        _set_cell_bg(data_row.cells[j], 'F5F8FA')
        p_dp = data_row.cells[j].add_paragraph()
        r_dp = p_dp.add_run(ph)
        r_dp.font.size = Pt(9); r_dp.font.name = 'Calibri'
        p_dp.paragraph_format.space_before = Pt(3)
        p_dp.paragraph_format.space_after = Pt(3)
    
    doc.add_paragraph()
    
    # Totals table
    totals_table = doc.add_table(rows=3, cols=2)
    totals_table.style = 'Table Grid'
    totals_table.alignment = WD_TABLE_ALIGNMENT.RIGHT
    
    totals_data = [
        ('Razem netto:', '[TOTAL_NETTO] PLN', 'FFFFFF', False),
        ('VAT [VAT_RATE]%:', '[VAT_AMOUNT] PLN', 'FFFFFF', False),
        ('RAZEM BRUTTO:', '[TOTAL_BRUTTO] PLN', '0D2137', True),
    ]
    for i, (lbl, val, bg, is_bold) in enumerate(totals_data):
        _set_cell_bg(totals_table.rows[i].cells[0], bg)
        _set_cell_bg(totals_table.rows[i].cells[1], bg)
        
        pl_t = totals_table.rows[i].cells[0].add_paragraph()
        rl_t = pl_t.add_run(lbl)
        rl_t.font.size = Pt(9.5); rl_t.font.bold = is_bold; rl_t.font.name = 'Calibri'
        rl_t.font.color.rgb = RGBColor(255,255,255) if is_bold else RGBColor(0x33,0x33,0x33)
        
        pv_t = totals_table.rows[i].cells[1].add_paragraph()
        pv_t.alignment = WD_ALIGN_PARAGRAPH.RIGHT
        rv_t = pv_t.add_run(val)
        rv_t.font.size = Pt(9.5 if not is_bold else 11)
        rv_t.font.bold = is_bold; rv_t.font.name = 'Calibri'
        rv_t.font.color.rgb = RGBColor(255,255,255) if is_bold else RGBColor(0x0D,0x21,0x37)
    
    doc.add_paragraph()
    add_styled_heading(doc, 'WARUNKI OFERTY', level=2, color='1565A0', size=12, space_before=6)
    
    terms_table = doc.add_table(rows=4, cols=2)
    terms_table.style = 'Table Grid'
    terms_labels = [
        ('Warunki płatności:', '[PAYMENT_TERMS]'),
        ('Czas dostawy:', '[DELIVERY_TIME]'),
        ('Gwarancja:', '[WARRANTY]'),
        ('Oferta ważna do:', '[VALID_UNTIL]'),
    ]
    for i, (lbl, val) in enumerate(terms_labels):
        terms_table.rows[i].cells[0].text = lbl
        terms_table.rows[i].cells[1].text = val
        _set_cell_bg(terms_table.rows[i].cells[0], 'F0F4F8')
        for cell in terms_table.rows[i].cells:
            for p in cell.paragraphs:
                p.paragraph_format.space_before = Pt(2)
                p.paragraph_format.space_after = Pt(2)
                for run in p.runs:
                    run.font.size = Pt(9.5)
                    run.font.name = 'Calibri'
        terms_table.rows[i].cells[0].paragraphs[0].runs[0].font.bold = True
        terms_table.rows[i].cells[0].paragraphs[0].runs[0].font.color.rgb = RGBColor(0x0D, 0x21, 0x37)
    
    doc.add_paragraph()
    p_notes = doc.add_paragraph('* [NOTES]')
    p_notes.paragraph_format.space_before = Pt(2)
    p_notes.paragraph_format.space_after = Pt(12)
    p_notes.runs[0].font.size = Pt(8)
    p_notes.runs[0].font.name = 'Calibri'
    p_notes.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    # Signature section
    add_styled_heading(doc, 'PODPISY', level=2, color='1565A0', size=12, space_before=8)
    sig_table = doc.add_table(rows=1, cols=2)
    sig_table.style = 'Table Grid'
    
    sig_labels = ['Pieczątka i podpis LaudaMed', 'Pieczątka i podpis Zamawiającego']
    for j, sl in enumerate(sig_labels):
        sig_cell = sig_table.rows[0].cells[j]
        sig_para = sig_cell.add_paragraph()
        sig_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        sig_para.paragraph_format.space_before = Pt(2)
        sig_run = sig_para.add_run(sl)
        sig_run.font.size = Pt(8); sig_run.font.name = 'Calibri'
        sig_run.font.color.rgb = RGBColor(0x66, 0x66, 0x66)
        
        blank = sig_cell.add_paragraph()
        blank.paragraph_format.space_before = Pt(30)
        blank.paragraph_format.space_after = Pt(2)
        
        date_para = sig_cell.add_paragraph()
        date_para.paragraph_format.space_before = Pt(2)
        date_para.paragraph_format.space_after = Pt(2)
        date_run = date_para.add_run('Data: _______________')
        date_run.font.size = Pt(8.5); date_run.font.name = 'Calibri'
        date_run.font.color.rgb = RGBColor(0x33, 0x33, 0x33)
    
    # Footer note
    doc.add_paragraph()
    p_foot = doc.add_paragraph(
        f"{COMPANY['name']}  |  {COMPANY['address']}  |  "
        f"Tel: {COMPANY['tel']}  |  {COMPANY['email']}  |  {COMPANY['web']}"
    )
    p_foot.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_foot.paragraph_format.space_before = Pt(8)
    p_foot.runs[0].font.size = Pt(7.5)
    p_foot.runs[0].font.name = 'Calibri'
    p_foot.runs[0].font.color.rgb = RGBColor(0x66, 0x66, 0x66)
    
    doc.save(output_path)
    print(f"[OK] DOCX template saved: {output_path}")
    return output_path

# ══════════════════════════════════════════════════════════════════════════════
# DEMO DATA — ViV60
# ══════════════════════════════════════════════════════════════════════════════

DEMO_OFFER_DATA = {
    "offer_number": "LM/2026/04/001",
    "date": "2026-04-02",
    "valid_until": "2026-05-02",
    "client_name": "Przychodnia POZ XYZ",
    "client_address": "ul. Przykładowa 1, 00-001 Warszawa",

    "product_name": "Aparat USG ViV60",
    "product_model": "ViV60",
    "product_category": "Ultrasonografia",
    "manufacturer": "ZONCARE",
    "cat_key": "USG",

    "product_image": "/home/user/workspace/extracted_images/img-008.png",
    "hero_image": "/home/user/workspace/extracted_images/img-008.png",

    "description": (
        "ViV60 to profesjonalny aparat USG klasy premium do diagnostyki w gabinecie "
        "i pracowni USG. Pokrywa pełen zakres badań w POZ: jama brzuszna, naczynia, "
        "tarczyca, ginekologia, MSK. Zaawansowane AI: Auto OB, Auto Follicle, HR Flow, "
        "zMicroFlow. Transfer przez WiFi, Bluetooth, DICOM 3.0."
    ),
    "benefits": [
        "Pełen zakres badań USG w POZ: jama brzuszna, naczynia, tarczyca, ginekologia, MSK",
        "AI: Auto OB, Auto EF, Auto IMT, Auto HRI, zMicroFlow — 14 modułów AI",
        "High-Resolution Blood Flow Imaging (HR Flow) — wykrywanie mikrokrążenia",
        "Integracja z HIS/PACS przez DICOM 3.0",
        "Zonnet Imaging — zdalna konsultacja ekspercka w czasie rzeczywistym",
        "3D/4D z ZLive, Elastografia, Panoramic Imaging (EFOV)",
    ],

    "specs": [
        ("Klasa", "Premium, stacjonarna, Color Doppler"),
        ("Platforma", "uSeed (CPU + GPU + FPGA)"),
        ("Monitor główny", "21,5\" LCD HD, obrót ±90°, pochył 0-90°"),
        ("Panel dotykowy", "13,3\" pojemnościowy, kąt 55°"),
        ("Porty sond", "4 aktywne + 1 port CW; podgrzewacz żelu 37/40°C"),
        ("Głowica convex", "3C6CD: 1.0–7.5 MHz | 128 elem."),
        ("Głowica liniowa", "7L4CD: 1.5–16.0 MHz | 128 elem."),
        ("Tryby obrazowania", "B, M, B/M, 2B/4B, CFM, PDI, PW, CW, Elastografia, 3D/4D"),
        ("AI funkcje", "Auto OB, Auto Follicle, HR Flow, Auto IMT, Auto EF, Auto HRI"),
        ("Transfery", "WiFi, Bluetooth, QR kod, Zonnet Imaging"),
        ("Interfejsy", "LAN, USB 3.0 x4, USB 2.0 x2, HDMI, DVI, VGA, DICOM 3.0"),
        ("Gwarancja", "24 miesiące"),
    ],

    "detailed_specs": [
        ("Klasa", "Premium, stacjonarna, Color Doppler"),
        ("Platforma", "uSeed (CPU + GPU + FPGA)"),
        ("Monitor główny", "21,5\" LCD HD, obrót ±90°, pochył 0-90°, wys. 235 mm"),
        ("Panel dotykowy", "13,3\" pojemnościowy, kąt 55°, wys. 0-160 mm"),
        ("Porty sond", "4 aktywne + 1 port sondy ołówkowej; podgrzewacz żelu 37/40°C"),
        ("Głowica convex", "3C6CD: 1.0–7.5 MHz | 128 elem. (centrum 3,5 MHz)"),
        ("Głowica liniowa", "7L4CD: 1.5–16.0 MHz | 128 elem. (centrum 7,5 MHz)"),
        ("Tryby obrazowania", "B, M, B/M, 2B/4B, CFM, PDI, PW, CW, Elastografia, 3D/4D"),
        ("Technologie obrazu", "Pure Harmonic, SCI, SRI/zonClear, zMicroFlow, EFOV, TDI"),
        ("AI funkcje", "Auto OB, Auto Follicle, HR Flow, Auto IMT, Auto EF, Auto HRI, zMicroFlow"),
        ("Transfery", "WiFi, Bluetooth, QR kod, Zonnet Imaging"),
        ("Interfejsy", "LAN, USB 3.0 x4, USB 2.0 x2, HDMI, DVI, VGA, RS-232, DICOM 3.0"),
        ("Wymiary systemu", "962 x 591 x 1305-1708 mm"),
        ("Masa systemu", "~100 kg"),
        ("System operacyjny", "Linux"),
        ("Gwarancja", "24 miesiące"),
    ],

    "extra_features": [
        "W zestawie: głowica convex 3C6CD + liniowa 7L4CD",
        "4 aktywne porty sond + 1 port sondy ołówkowej CW",
        "Podgrzewacz żelu zintegrowany w panelu głównym (37/40°C)",
        "Ergonomiczny panel dotykowy 13,3\" z regulacją kąta 55°",
        "Monitor obrotowy ±90°, pochylany 0-90°, wysokość regulowana 235 mm",
        "Zonnet Imaging — zdalna konsultacja ekspercka w czasie rzeczywistym",
        "Instalacja, konfiguracja i szkolenie w cenie",
    ],

    "accessories": [
        ("Głowica endowaginalna 6E1CDS", "na zapytanie"),
        ("Głowica 4D volume 4V4CD", "na zapytanie"),
        ("Głowica liniowa HD 10L5CD", "na zapytanie"),
        ("Głowica kardio-pediatryczna 5S1CD", "na zapytanie"),
        ("Wózek mobilny z koszem", "2 500 PLN netto"),
        ("Drukarka USG Sony UP-X898MD", "na zapytanie"),
        ("Przedłużona gwarancja (36 mies.)", "na zapytanie"),
    ],

    "pricing": [
        {
            "model": "ViV60 (konfiguracja standardowa)",
            "description": "2 głowice: convex 3C6CD + liniowa 7L4CD",
            "qty": 1,
            "unit_price_netto": 108200,
        },
    ],

    "vat_rate": 0.08,
    "payment_terms": "30 dni od daty dostawy",
    "delivery_time": "4-6 tygodni od zamówienia",
    "warranty": "24 miesiące (serwis gwarancyjny na miejscu)",
    "certs": ["CE", "MDR", "ISO 13485"],

    "notes": [
        "Ceny netto, należy doliczyć VAT 8%.",
        "Instalacja, konfiguracja i szkolenie personelu w cenie dostawy.",
        "Ceny ważne do daty określonej w nagłówku oferty.",
    ],
}

# ══════════════════════════════════════════════════════════════════════════════
# CLI ENTRY POINT
# ══════════════════════════════════════════════════════════════════════════════

def main():
    parser = argparse.ArgumentParser(description='LaudaMed Offer Generator')
    parser.add_argument('--data', metavar='FILE', help='JSON file with offer data → generate PDF')
    parser.add_argument('--template', action='store_true', help='Generate DOCX template')
    parser.add_argument('--demo', action='store_true', help='Generate demo PDF + DOCX for ViV60')
    parser.add_argument('--output', metavar='PATH', help='Output file path (optional)')
    args = parser.parse_args()
    
    os.makedirs(OUTPUT_DIR, exist_ok=True)
    
    if args.demo:
        # Generate both PDF and DOCX demo
        print("=== LaudaMed Offer Generator — DEMO MODE ===")
        pdf_path = args.output or os.path.join(OUTPUT_DIR, 'oferta_ViV60_LM-2026-04-001.pdf')
        docx_path = os.path.join(OUTPUT_DIR, 'oferta_template.docx')
        
        print(f"\n[PDF] Generating {pdf_path}")
        generate_pdf(DEMO_OFFER_DATA, pdf_path)
        
        print(f"\n[DOCX] Generating {docx_path}")
        generate_docx_template(docx_path)
        
        print("\n=== Done! ===")
        print(f"PDF:  {pdf_path}")
        print(f"DOCX: {docx_path}")
    
    elif args.data:
        with open(args.data, 'r', encoding='utf-8') as f:
            offer_data = json.load(f)
        
        name_slug = offer_data.get('product_model', 'offer').replace(' ', '_')
        num_slug = offer_data.get('offer_number', '').replace('/', '-')
        pdf_path = args.output or os.path.join(OUTPUT_DIR, f'oferta_{name_slug}_{num_slug}.pdf')
        
        print(f"[PDF] Generating {pdf_path}")
        generate_pdf(offer_data, pdf_path)
    
    elif args.template:
        docx_path = args.output or os.path.join(OUTPUT_DIR, 'oferta_template.docx')
        print(f"[DOCX] Generating {docx_path}")
        generate_docx_template(docx_path)
    
    else:
        parser.print_help()
        print("\nExample usage:")
        print("  python3 generate_offer.py --demo")
        print("  python3 generate_offer.py --data offer_data.json")
        print("  python3 generate_offer.py --template")

if __name__ == '__main__':
    main()
